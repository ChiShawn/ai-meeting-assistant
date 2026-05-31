#!/usr/bin/env python3
"""Portfolio-safe FastAPI backend for AI 會議記錄助手.

The production version used a larger domain-specific DOCX filling pipeline. This
public version keeps the same architecture and API surface while avoiding local
private deployment details.
"""

from __future__ import annotations

import asyncio
import io
import os
import tempfile
import uuid
from pathlib import Path
from typing import Dict, List, Optional

import httpx
from fastapi import BackgroundTasks, FastAPI, File, Form, HTTPException, UploadFile
from fastapi.middleware.cors import CORSMiddleware
from fastapi.responses import HTMLResponse, Response
from pydantic import BaseModel

try:
    from docx import Document
except Exception:  # pragma: no cover - optional dependency at import time
    Document = None


VLLM_API_URL = os.getenv("VLLM_API_URL", "http://127.0.0.1:18080/v1/chat/completions")
VLLM_MODEL_NAME = os.getenv("VLLM_MODEL_NAME", "Qwen/Qwen3-0.6B")
REMOTE_ASR_URL = os.getenv("REMOTE_ASR_URL", "http://127.0.0.1:18080/asr")
REQUEST_TIMEOUT = float(os.getenv("REQUEST_TIMEOUT", "10800"))
ALLOW_ORIGINS = os.getenv("ALLOW_ORIGINS", "*").split(",")

app = FastAPI(title="AI 會議記錄助手", version="1.0.0")
app.add_middleware(
    CORSMiddleware,
    allow_origins=ALLOW_ORIGINS,
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)

jobs: Dict[str, dict] = {}
corrections: Dict[str, str] = {}


class SummarizeRequest(BaseModel):
    text_to_summarize: str
    summary_length: Optional[str] = "medium"


class SummarizeResponse(BaseModel):
    summary_text: str
    error: Optional[str] = None


class ToDoRequest(BaseModel):
    text_to_analyze: str


class ToDoResponse(BaseModel):
    todo_list: str
    error: Optional[str] = None


class MinutesItemsRequest(BaseModel):
    text_to_analyze: str
    summary_text: Optional[str] = ""


class MinutesItem(BaseModel):
    section: str = ""
    title: str = ""
    content: str


class MinutesItemsResponse(BaseModel):
    items: List[MinutesItem]
    error: Optional[str] = None


class CorrectionRequest(BaseModel):
    wrong_word: str
    correct_word: str


def hms(seconds: float) -> str:
    seconds = max(0, int(round(seconds)))
    return f"{seconds // 3600:02d}:{(seconds % 3600) // 60:02d}:{seconds % 60:02d}"


def normalize_segments(raw_segments: list) -> list:
    normalized = []
    for seg in raw_segments or []:
        if not isinstance(seg, dict):
            continue
        start = seg.get("start", seg.get("start_time", 0))
        end = seg.get("end", seg.get("end_time", start))
        text = str(seg.get("text", seg.get("sentence", ""))).strip()
        try:
            start_f = float(start)
            end_f = float(end)
        except Exception:
            start_f = end_f = 0.0
        for wrong, correct in corrections.items():
            text = text.replace(wrong, correct)
        normalized.append({"start": hms(start_f), "end": hms(end_f), "text": text})
    return normalized


def transcript_text(segments: list, fallback: str = "") -> str:
    if segments:
        return "".join(str(seg.get("text", "")) for seg in segments).strip()
    return fallback.strip()


async def call_llm(system_prompt: str, user_prompt: str, max_tokens: int = 1024) -> str:
    payload = {
        "model": VLLM_MODEL_NAME,
        "messages": [
            {"role": "system", "content": system_prompt},
            {"role": "user", "content": user_prompt},
        ],
        "max_tokens": max_tokens,
        "temperature": 0.2,
        "stream": False,
        "chat_template_kwargs": {"enable_thinking": False},
    }
    async with httpx.AsyncClient(timeout=REQUEST_TIMEOUT) as client:
        response = await client.post(VLLM_API_URL, json=payload)
        response.raise_for_status()
    data = response.json()
    return data["choices"][0]["message"]["content"].strip()


async def run_remote_asr(file_bytes: bytes, filename: str, language: str = "zh") -> dict:
    files = {"file": (filename, file_bytes, "application/octet-stream")}
    data = {"language": language, "diarization": "0"}
    async with httpx.AsyncClient(timeout=REQUEST_TIMEOUT) as client:
        response = await client.post(REMOTE_ASR_URL, files=files, data=data)
        response.raise_for_status()
    payload = response.json()
    segments = normalize_segments(payload.get("segments", []))
    return {
        "text": transcript_text(segments, payload.get("text", "")),
        "segments": segments,
        "raw_backend": payload.get("model", "remote-asr"),
    }


async def transcribe_worker(task_id: str, file_bytes: bytes, filename: str, language: str) -> None:
    jobs[task_id].update(status="running", progress=20)
    try:
        result = await run_remote_asr(file_bytes, filename, language)
        jobs[task_id].update(status="done", progress=100, result=result)
    except Exception as exc:
        jobs[task_id].update(status="error", progress=100, error=str(exc))


async def delayed_job_cleanup(task_id: str, delay_seconds: int = 600) -> None:
    await asyncio.sleep(delay_seconds)
    jobs.pop(task_id, None)


@app.get("/health")
async def health():
    return {"ok": True, "jobs": len(jobs), "remote_asr_url": bool(REMOTE_ASR_URL), "vllm_api_url": bool(VLLM_API_URL)}


@app.get("/", response_class=HTMLResponse)
async def index():
    return "<h1>AI 會議記錄助手</h1><p>Serve frontend/index.html through Apache/PHP for the full UI.</p>"


@app.post("/transcribe_async")
async def transcribe_async(
    background_tasks: BackgroundTasks,
    file: UploadFile = File(...),
    language: str = Form("zh"),
    model_size: str = Form("remote"),
):
    file_bytes = await file.read()
    if not file_bytes:
        raise HTTPException(status_code=400, detail="empty upload")
    task_id = str(uuid.uuid4())
    jobs[task_id] = {"status": "queued", "progress": 0, "filename": file.filename}
    background_tasks.add_task(transcribe_worker, task_id, file_bytes, file.filename or "upload.audio", language)
    return {"task_id": task_id, "status": "queued"}


@app.get("/task/{task_id}")
async def get_task(task_id: str, background_tasks: BackgroundTasks):
    job = jobs.get(task_id)
    if not job:
        raise HTTPException(status_code=404, detail="task not found")
    if job.get("status") in {"done", "error"}:
        background_tasks.add_task(delayed_job_cleanup, task_id)
    return job


@app.delete("/task/{task_id}")
async def delete_task(task_id: str):
    jobs.pop(task_id, None)
    return {"ok": True}


@app.post("/summarize", response_model=SummarizeResponse)
async def summarize(req: SummarizeRequest):
    if not req.text_to_summarize.strip():
        return SummarizeResponse(summary_text="", error="empty text")
    try:
        content = await call_llm(
            "你是專業會議記錄整理助手，請用繁體中文輸出清楚、精簡、可追蹤的會議摘要。",
            f"請整理以下逐字稿，產生重點摘要。摘要長度偏好：{req.summary_length}\n\n{req.text_to_summarize}",
            max_tokens=1200,
        )
        return SummarizeResponse(summary_text=content)
    except Exception as exc:
        return SummarizeResponse(summary_text=f"摘要產生失敗：{exc}", error=str(exc))


@app.post("/todo", response_model=ToDoResponse)
async def todo(req: ToDoRequest):
    if not req.text_to_analyze.strip():
        return ToDoResponse(todo_list="", error="empty text")
    try:
        content = await call_llm(
            "你是專案管理助理，請只擷取具體待辦事項、負責人、期限與追蹤重點。",
            f"請從以下會議逐字稿擷取待辦事項，使用條列式繁體中文：\n\n{req.text_to_analyze}",
            max_tokens=900,
        )
        return ToDoResponse(todo_list=content)
    except Exception as exc:
        return ToDoResponse(todo_list=f"待辦提取失敗：{exc}", error=str(exc))


@app.post("/minutes_items", response_model=MinutesItemsResponse)
async def minutes_items(req: MinutesItemsRequest):
    try:
        prompt = (
            "請把會議內容整理成可放入 Word 會議記錄模板的條列。"
            "每一點需包含建議章節、標題與內容。請用 JSON array，欄位為 section/title/content。\n\n"
            f"摘要參考：{req.summary_text or ''}\n\n逐字稿：{req.text_to_analyze}"
        )
        raw = await call_llm("你輸出嚴格 JSON，不要 markdown。", prompt, max_tokens=1200)
        import json

        parsed = json.loads(raw)
        items = [MinutesItem(**item) for item in parsed if item.get("content")]
        return MinutesItemsResponse(items=items)
    except Exception as exc:
        fallback = req.summary_text or req.text_to_analyze[:1000]
        return MinutesItemsResponse(items=[MinutesItem(section="摘要", title="會議重點", content=fallback)], error=str(exc))


@app.get("/corrections")
async def list_corrections():
    return {"simple_replacements": [{"wrong_word": k, "correct_word": v} for k, v in corrections.items()]}


@app.post("/corrections/add")
async def add_correction(req: CorrectionRequest):
    wrong = req.wrong_word.strip()
    correct = req.correct_word.strip()
    if not wrong or not correct:
        raise HTTPException(status_code=400, detail="wrong_word and correct_word are required")
    corrections[wrong] = correct
    return {"ok": True, "wrong_word": wrong, "correct_word": correct}


@app.post("/upload_probe")
async def upload_probe(file: UploadFile = File(...)):
    data = await file.read()
    return {"ok": True, "filename": file.filename, "size": len(data), "tmp_readable": True}


@app.post("/generate_report_docx")
async def generate_report_docx(
    template_file: UploadFile = File(...),
    summary_text: str = Form(""),
    todo_text: str = Form(""),
    source_text: str = Form(""),
    minutes_items_json: str = Form(""),
):
    if Document is None:
        raise HTTPException(status_code=500, detail="python-docx is not installed")

    template_bytes = await template_file.read()
    try:
        doc = Document(io.BytesIO(template_bytes))
    except Exception:
        doc = Document()
        doc.add_heading("AI 會議記錄", level=1)

    replacements = {
        "{{摘要}}": summary_text,
        "{{summary}}": summary_text,
        "{{待辦}}": todo_text,
        "{{todo}}": todo_text,
        "{{逐字稿}}": source_text,
        "{{transcript}}": source_text,
    }
    replaced = False
    for paragraph in doc.paragraphs:
        for key, value in replacements.items():
            if key in paragraph.text:
                paragraph.text = paragraph.text.replace(key, value or "")
                replaced = True

    if not replaced:
        doc.add_heading("會議摘要", level=2)
        doc.add_paragraph(summary_text or "（無摘要）")
        doc.add_heading("待辦事項", level=2)
        doc.add_paragraph(todo_text or "（無待辦）")

    out = io.BytesIO()
    doc.save(out)
    out.seek(0)
    return Response(
        content=out.read(),
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        headers={"Content-Disposition": 'attachment; filename="meeting_report.docx"'},
    )
